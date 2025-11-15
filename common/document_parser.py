"""
文档解析器模块
支持PDF和Word文档的文本提取和预处理
"""
import os
import re
from abc import ABC, abstractmethod
from typing import List, Optional
import logging

# 设置日志
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentParser(ABC):
    """文档解析器基类"""

    @abstractmethod
    def extract_text(self, file_path: str) -> str:
        """
        提取文档中的文本内容

        Args:
            file_path: 文档文件路径

        Returns:
            提取的文本内容
        """
        pass

    def _clean_page_markers(self, text: str) -> str:
        """
        清理页码、页眉页脚等标记

        Args:
            text: 原始文本

        Returns:
            清理后的文本
        """
        # 移除常见的页码模式
        patterns = [
            r'第\s*\d+\s*页',  # 中文页码
            r'Page\s*\d+',      # 英文页码
            r'\d+\s*/\s*\d+',   # 页码格式 1/10
            r'^\d+$',           # 独立的数字行
            r'-\s*\d+\s*-',     # - 1 - 格式
        ]

        for pattern in patterns:
            text = re.sub(pattern, '', text, flags=re.IGNORECASE | re.MULTILINE)

        # 移除多余的空行
        text = re.sub(r'\n{3,}', '\n\n', text)

        return text.strip()


class PDFParser(DocumentParser):
    """PDF文档解析器"""

    def extract_text(self, file_path: str) -> str:
        """
        从PDF文件中提取文本

        Args:
            file_path: PDF文件路径

        Returns:
            提取的文本内容
        """
        try:
            import fitz  # PyMuPDF

            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            text_blocks = []

            # 打开PDF文档
            with fitz.open(file_path) as doc:
                logger.info(f"开始解析PDF文档，共 {len(doc)} 页")

                for page_num, page in enumerate(doc, 1):
                    # 提取当前页的文本
                    text = page.get_text()

                    # 清理页面标记
                    text = self._clean_page_markers(text)

                    # 移除空白页
                    if text.strip():
                        text_blocks.append(text)
                        logger.debug(f"已处理第 {page_num} 页")

            # 合并所有页面的文本
            full_text = "\n\n".join(text_blocks)

            logger.info(f"PDF解析完成，提取文本长度: {len(full_text)} 字符")
            return full_text

        except ImportError:
            raise ImportError("请安装PyMuPDF库: pip install PyMuPDF")
        except Exception as e:
            logger.error(f"PDF解析失败: {str(e)}")
            raise

    def extract_text_with_structure(self, file_path: str) -> List[dict]:
        """
        提取PDF文本并保留结构信息（章节、段落等）

        Args:
            file_path: PDF文件路径

        Returns:
            包含结构信息的文本块列表
        """
        try:
            import fitz

            structured_content = []

            with fitz.open(file_path) as doc:
                for page_num, page in enumerate(doc, 1):
                    # 获取页面的文本块
                    blocks = page.get_text("blocks")

                    for block in blocks:
                        if block[4].strip():  # block[4]是文本内容
                            structured_content.append({
                                'page': page_num,
                                'text': block[4].strip(),
                                'bbox': block[:4],  # 边界框坐标
                                'type': 'text'
                            })

            return structured_content

        except Exception as e:
            logger.error(f"结构化提取失败: {str(e)}")
            return []


class WordParser(DocumentParser):
    """Word文档解析器"""

    def extract_text(self, file_path: str) -> str:
        """
        从Word文档中提取文本

        Args:
            file_path: Word文档路径

        Returns:
            提取的文本内容
        """
        try:
            from docx import Document

            if not os.path.exists(file_path):
                raise FileNotFoundError(f"文件不存在: {file_path}")

            doc = Document(file_path)
            text_blocks = []

            logger.info(f"开始解析Word文档")

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    # 识别标题样式
                    if para.style and 'Heading' in para.style.name:
                        text = f"\n## {text}\n"  # 添加标题标记
                    text_blocks.append(text)

            # 提取表格内容
            for table in doc.tables:
                table_text = self._extract_table_text(table)
                if table_text:
                    text_blocks.append(table_text)

            # 合并所有文本
            full_text = "\n\n".join(text_blocks)

            # 清理页面标记
            full_text = self._clean_page_markers(full_text)

            logger.info(f"Word解析完成，提取文本长度: {len(full_text)} 字符")
            return full_text

        except ImportError:
            raise ImportError("请安装python-docx库: pip install python-docx")
        except Exception as e:
            logger.error(f"Word文档解析失败: {str(e)}")
            raise

    def _extract_table_text(self, table) -> str:
        """
        从表格中提取文本

        Args:
            table: Word表格对象

        Returns:
            表格文本内容
        """
        rows = []

        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cells.append(cell_text)

            if cells:
                rows.append(" | ".join(cells))

        if rows:
            # 添加表格分隔符
            header_separator = "-" * 40
            return f"{header_separator}\n" + "\n".join(rows) + f"\n{header_separator}"

        return ""

    def extract_with_formatting(self, file_path: str) -> List[dict]:
        """
        提取Word文档内容并保留格式信息

        Args:
            file_path: Word文档路径

        Returns:
            包含格式信息的内容列表
        """
        try:
            from docx import Document

            doc = Document(file_path)
            formatted_content = []

            for para in doc.paragraphs:
                if para.text.strip():
                    formatted_content.append({
                        'text': para.text.strip(),
                        'style': para.style.name if para.style else None,
                        'is_heading': 'Heading' in (para.style.name if para.style else ''),
                        'type': 'paragraph'
                    })

            for table_idx, table in enumerate(doc.tables):
                formatted_content.append({
                    'text': self._extract_table_text(table),
                    'type': 'table',
                    'table_index': table_idx
                })

            return formatted_content

        except Exception as e:
            logger.error(f"格式化提取失败: {str(e)}")
            return []


class ParserFactory:
    """解析器工厂类"""

    # 支持的文件扩展名映射
    SUPPORTED_EXTENSIONS = {
        '.pdf': PDFParser,
        '.docx': WordParser,
        '.doc': WordParser,  # 注意：.doc格式可能需要额外处理
    }

    @classmethod
    def get_parser(cls, file_path: str) -> DocumentParser:
        """
        根据文件类型获取相应的解析器

        Args:
            file_path: 文件路径

        Returns:
            对应的文档解析器实例

        Raises:
            ValueError: 不支持的文件类型
        """
        # 获取文件扩展名
        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        # 检查是否支持该文件类型
        if ext not in cls.SUPPORTED_EXTENSIONS:
            supported = ', '.join(cls.SUPPORTED_EXTENSIONS.keys())
            raise ValueError(f"不支持的文件类型: {ext}。支持的类型: {supported}")

        # 返回对应的解析器实例
        parser_class = cls.SUPPORTED_EXTENSIONS[ext]
        return parser_class()

    @classmethod
    def is_supported(cls, file_path: str) -> bool:
        """
        检查文件类型是否被支持

        Args:
            file_path: 文件路径

        Returns:
            是否支持该文件类型
        """
        _, ext = os.path.splitext(file_path)
        return ext.lower() in cls.SUPPORTED_EXTENSIONS


# 便捷函数
def extract_document_text(file_path: str) -> str:
    """
    便捷函数：直接从文档中提取文本

    Args:
        file_path: 文档路径

    Returns:
        提取的文本内容
    """
    parser = ParserFactory.get_parser(file_path)
    return parser.extract_text(file_path)


if __name__ == "__main__":
    # 测试代码
    import sys

    if len(sys.argv) < 2:
        print("用法: python document_parser.py <文档路径>")
        sys.exit(1)

    file_path = sys.argv[1]

    try:
        # 获取解析器
        parser = ParserFactory.get_parser(file_path)

        # 提取文本
        text = parser.extract_text(file_path)

        # 显示结果
        print(f"文档类型: {type(parser).__name__}")
        print(f"提取文本长度: {len(text)} 字符")
        print("\n前500个字符:")
        print("-" * 50)
        print(text[:500])
        print("-" * 50)

    except Exception as e:
        print(f"错误: {e}")
        sys.exit(1)